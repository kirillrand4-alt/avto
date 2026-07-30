# -*- coding: utf-8 -*-
"""Вложения закупок ЕИС: техзадания, извещения, протоколы — там, где живёт свободный текст.

Зачем. Перенос приёма Tender.pro на ЕИС проверен и **в теле карточки не работает**: блок
«Дополнительная информация» отсутствует у 43 карточек из 48 и содержит прочерк у пяти
(`engineers-lens/centro/eis/ZAMER-svobodnyy-tekst.md`). Но вывод там сужен до объекта замера:
свободного текста нет **в карточке**, а у ЕИС он живёт во **вложениях**. Живой образец уже
есть: у Газпром нефтехим Салавата в документе стояло «Инициатор мероприятия: Начальник УЭПБ и
ОТ В. А. Кузнецов», а в листе согласования «Главный механик (первичной переработки НПЗ)
Э. В. Ваганов».

Путь до файла найден в коде уже скачанных карточек, а не угадан:
    карточка → href «/epz/order/notice/notice223/documents.html?purchaseNoticeNumber=…&noticeGuid=…»
    страница документов → href «https://zakupki.gov.ru/223/filestore/public/1.0/download/fz223/file.html?uid=…»

Заслоны, каждый по конкретной ошибке прошлых прогонов:

- **Текст из PDF берётся библиотекой, а не разбором сырых байт.** В прошлом массовом обогащении
  регулярка по десяти цифрам собрала 3 815 ложных телефонов из координат векторной графики
  (`M8892 7.44042 1`). Библиотека отдаёт текстовый слой, координат в нём нет.
- **Порог распознавания по длине, а не «если пусто»:** меньше 200 знаков на файл — помечаем
  сканом. У соседней сессии 39 из 74 PDF оказались сканами с обрывком текстового слоя из
  колонтитула, и порог «если пусто» их пропустил бы.
- **Расширение врёт.** Тип определяется по сигнатуре первых байт (`%PDF`, `PK`, `\\xd0\\xcf`),
  потому что из 69 «доков» у соседей 33 оказались HTML внутри.
- **Номер принимается только рядом со словом связи** (тел/моб/контакт/сот) и **не принимается,
  если совпал с ИНН или ОГРН, напечатанным в том же файле**.

Шаги (каждый отдельно, результат на диске, повтор не переделывает сделанное):
    python3 vlozheniya.py --doki      # карточки → страницы документов
    python3 vlozheniya.py --fajly     # страницы документов → скачанные вложения
    python3 vlozheniya.py --tekst     # вложения → текст + пометка «скан»
    python3 vlozheniya.py --lica      # текст → люди, провайдером
"""
import csv
import glob
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import threading
import time
import zipfile
from concurrent.futures import ThreadPoolExecutor

BAZA = os.path.dirname(os.path.abspath(__file__))
KLIENT = os.path.join(BAZA, 'server', 'run_on_server.py')
DROP = os.path.join(BAZA, 'server', 'drop_client.sh')
# Рабочая папка кэша. Была прибита к пути scratchpad конкретной сессии
# (`/tmp/claude-0/…/520847fd-…/scratchpad`) — в следующей сессии этой папки уже нет, шаги видят
# ноль файлов и это выглядит как «ничего не скачано». Сами файлы живут на дропе, локально они
# только кэш, поэтому путь переопределяется переменной EIS_RAB.
RAB = os.environ.get('EIS_RAB', '/home/user/work/eis')
KART = os.path.join(RAB, 'eis')
DOKI = os.path.join(RAB, 'eis_doki')
FAJLY = os.path.join(RAB, 'eis_fajly')
OUT = os.path.join(BAZA, 'engineers-lens', 'centro', 'eis')

SSYLKA_DOK = re.compile(r'href="(/epz/order/notice/notice223/documents\.html\?[^"]+)"')
SSYLKA_FAJL = re.compile(r'href="(https://zakupki\.gov\.ru/223/filestore/public/1\.0/download/'
                         r'fz223/file\.html\?uid=[0-9A-F]+)"', re.I)
# Должность рядом с ФИО — то, ради чего всё делается.
#
# Оба шаблона расширены 30.07.2026 после контроля на эталонах. Прежние проваливали два случая
# из пяти, и оба — живые записи из наших же данных:
#   «Вопросы технического характера: **Александр Владимирович Лосев**, Заместитель главного
#    механика» — фамилия стоит ПОСЛЕДНЕЙ, а шаблон ждал её первой;
#   «Инициатор мероприятия: **Начальник УЭПБ и ОТ** В. А. Кузнецов» — должности-аббревиатуры
#    (УЭПБ, ОГМ, ОГЭ, УГЭ, КИПиА, ОТиПБ) в списке не было вовсе.
# Это ровно то, из-за чего первый замер вложений дал одну пару на 106 файлов: мерился не
# источник, а собственный шаблон.
DOLZH = re.compile(
    r'((?:главн\w+|ведущ\w+|старш\w+|заместител\w+\s+главн\w+|зам\.?\s*главн\w+|'
    r'и\.?\s*о\.?\s*главн\w+)?\s*'
    r'(?:инженер\w*|механик\w*|энергетик\w*|технолог\w*|метролог\w*)'
    r'(?:\s+(?:по\s+[\w\-]+|цеха|производства|участка|управления)){0,2}'
    r'|начальник\w*\s+(?:цеха|компрессорн\w+|отдела\s+глав\w+|управления\s+глав\w+|'
    r'производств\w+|участка|службы\s+\w+|[А-ЯЁ]{2,6}(?:\s+и\s+[А-ЯЁ]{2,4})?)'
    r'|(?:главн\w+|ведущ\w+)\s+специалист\w*'
    r'|технически\w+\s+(?:директор\w*|руководител\w*)'
    r'|руководител\w+\s+(?:технической\s+службы|службы\s+глав\w+))', re.I)
FIO = re.compile(r'([А-ЯЁ][а-яё\-]{2,}\s+[А-ЯЁ]\.\s?[А-ЯЁ]\.'          # Иванов И. И.
                 r'|[А-ЯЁ]\.\s?[А-ЯЁ]\.\s?[А-ЯЁ][а-яё\-]{2,}'          # И. И. Иванов
                 r'|[А-ЯЁ][а-яё\-]{2,}\s+[А-ЯЁ][а-яё\-]{2,}(?:вич|вна)\b'   # Иванов Иванович
                 r'|[А-ЯЁ][а-яё\-]{2,}\s+[А-ЯЁ][а-яё\-]{2,}(?:вич|вна)\s+'
                 r'[А-ЯЁ][а-яё\-]{2,}\b'                               # Иван Иванович Иванов
                 r'|[А-ЯЁ][а-яё\-]{2,}\s+[А-ЯЁ][а-яё\-]{2,}\s+'
                 r'[А-ЯЁ][а-яё\-]{2,}(?:вич|вна)\b)')                  # Иванов Иван Иванович
SVYAZ = re.compile(r'(?:тел|моб|сот|контакт|факс|т\.)\D{0,12}'
                   r'((?:\+7|\b8)?[\s(\-]*\d{3,5}[\s)\-]*\d{2,3}[\s\-]*\d{2}[\s\-]*\d{2})', re.I)
INN_OGRN = re.compile(r'\b(\d{10}|\d{12}|\d{13}|\d{15})\b')
# Предел на сохраняемый текст одного файла. Замер: прежние 2 000 000 резали 27 файлов, у которых
# настоящих знаков 257 млн против сохранённых 54 млн. Самый длинный файл корпуса — 12,6 млн.
PREDEL_TEKSTA = 20_000_000
# Предел страниц на OCR одного файла. Не «на всякий случай»: живой прогон встал на 36-мегабайтном
# `.rar` со сканами внутри и держал очередь больше двадцати минут — распознавание идёт секундами
# на страницу, а таких файлов в корпусе десятки. Упор в предел ПЕЧАТАЕТСЯ и помечается в строке
# (`ocr_obrezan`), потому что молчаливый предел — это ровно то, что мы весь день выкорчёвываем.
PREDEL_STRANIC_OCR = 60
# Общий бюджет страниц OCR на ОДИН файл корпуса. Предела на отдельный PDF мало: внутри архива
# таких PDF десятки, и 36-мегабайтный `.rar` держал очередь больше двадцати минут. Бюджет
# тратится всеми вложенными файлами вместе и обнуляется на каждом файле корпуса.
PREDEL_STRANIC_OCR_FAJL = 150
_bujet = {'stranic': PREDEL_STRANIC_OCR_FAJL, 'obrezano': 0}


def cifry(s):
    return re.sub(r'\D', '', s or '')


def podpis(b, p=None):
    """Тип по сигнатуре, а не по расширению: расширение врёт.

    Разделение `PK` на docx/xlsx и настоящий архив добавлено 30.07.2026, и это не тонкость.
    Прежде всё, что начинается с `PK`, считалось `zip/docx/xlsx` и уходило в `docx.Document()`;
    для архива с PDF внутри это падало, запасной путь собирал из архива только файлы `.xml`, а
    их там нет — текст выходил пустой, файл помечался сканом. Так «Документация.zip» на 553 КБ
    прошла как «пусто». Различать надо по содержимому архива: у docx внутри лежит
    `word/document.xml`, у xlsx — `xl/workbook.xml`, у настоящего архива — обычные файлы.
    """
    if b[:4] == b'%PDF':
        return 'pdf'
    if b[:6] == b'7z\xbc\xaf\x27\x1c':
        return '7z'
    if b[:4] == b'Rar!':
        return 'rar'
    if b[:2] == b'PK':
        if p:
            try:
                with zipfile.ZipFile(p) as z:
                    imena = z.namelist()
                if any(n.startswith('word/document') for n in imena):
                    return 'docx'
                if any(n.startswith('xl/workbook') for n in imena):
                    return 'xlsx'
                if any(n.startswith('content.xml') or n.startswith('ppt/') for n in imena):
                    return 'ooxml-иное'
                return 'архив'
            except Exception:  # noqa: BLE001  битый zip — пусть решает извлекатель
                return 'zip-битый'
        return 'zip/docx/xlsx'
    if b[:4] == b'\xd0\xcf\x11\xe0':
        return 'doc/xls'
    if b[:15].lower().lstrip().startswith((b'<html', b'<!doctype')):
        return 'html'
    if b[:5] == b'{\\rtf':
        return 'rtf'
    # Отказ ЕИС — отдельный тип, а не «иное». Хранилище отвечает JSON-ом
    # `{"message":"Ошибка доступа. Файл с uri FZ223/… не опубликован","status":"ERROR"}` на
    # 142 байта. Прежде он попадал в «иное», текст выходил пустой, и файл считался сканом:
    # 94 отказа из 292 файлов лежали в корзине «скан или пусто». Отказ и пустота — разные вещи:
    # пустой скан надо распознавать, а отказ надо перезапрашивать или признавать
    # неопубликованным. Пока они в одной корзине, ни того ни другого не видно.
    if b[:1] == b'{' and b'"status":"ERROR"' in b[:400]:
        return 'отказ'
    return 'иное'


def runner_many(zad, threads=6):
    p = subprocess.run([sys.executable, KLIENT, '--many', json.dumps(zad, ensure_ascii=False),
                        '--threads', str(threads)], capture_output=True, text=True, timeout=1800)
    m = re.search(r'\[.*\]', p.stdout, re.S)
    return json.loads(m.group(0)) if m else []


def skachat(imya, kuda):
    os.makedirs(kuda, exist_ok=True)
    subprocess.run(['bash', DROP, 'down', imya], cwd=kuda, capture_output=True, timeout=600)
    p = os.path.join(kuda, imya)
    return p if os.path.exists(p) else ''


def shag_doki(pachka=6):
    os.makedirs(DOKI, exist_ok=True)
    zad = []
    for p in sorted(glob.glob(os.path.join(KART, 'eis_c_*.html'))):
        n = re.search(r'eis_c_(\d+)', p).group(1)
        if os.path.exists(os.path.join(DOKI, f'eis_d_{n}.html')):
            continue
        h = open(p, encoding='utf-8', errors='replace').read()
        m = SSYLKA_DOK.search(h)
        if m:
            zad.append({'task': 'fetch_url',
                        'args': {'url': 'https://zakupki.gov.ru' + m.group(1).replace('&amp;', '&'),
                                 'insecure': True, 'name': f'eis_d_{n}.html'}})
    print(f'страниц документов к обходу: {len(zad)}', file=sys.stderr)
    for i in range(0, len(zad), pachka):
        for r in runner_many(zad[i:i + pachka], pachka):
            d = (r or {}).get('data') or {}
            if d.get('drop_name'):
                skachat(d['drop_name'], DOKI)
        print(f'  {min(i + pachka, len(zad))}/{len(zad)}', file=sys.stderr, flush=True)


# Порядок докачки по разряду документа. Прежний слепой предел `zad[:200]` резал очередь в том
# порядке, в каком она собралась, — по возрастанию номера закупки. Из-за этого до 2026 года
# очередь не доходила вовсе, а протоколы (которые нам не нужны) качались раньше техзаданий.
# Чем меньше число, тем раньше качается.
VES_RAZRYADA = [
    (1, re.compile(r'техническ\w*\s*(?:задани|част)|\bТЗ\b|техзадани', re.I)),
    (2, re.compile(r'документаци', re.I)),
    (3, re.compile(r'руководств|эксплуатац|паспорт', re.I)),
    (4, re.compile(r'приложени|специфик|опросн', re.I)),
    (5, re.compile(r'договор|контракт', re.I)),
    (6, re.compile(r'извещени|обоснован|НМЦ', re.I)),
    (8, re.compile(r'протокол|разъяснени|изменени|отмен', re.I)),
]


def ves(imya):
    for v, r in VES_RAZRYADA:
        if r.search(imya or ''):
            return v
    return 7


def shag_fajly(pachka=6, predel=200):
    os.makedirs(FAJLY, exist_ok=True)
    imena = imena_dokumentov()
    zad = []
    for p in sorted(glob.glob(os.path.join(DOKI, 'eis_d_*.html'))):
        n = re.search(r'eis_d_(\d+)', p).group(1)
        h = open(p, encoding='utf-8', errors='replace').read()
        for j, u in enumerate(dict.fromkeys(SSYLKA_FAJL.findall(h)), 1):
            imya = f'eis_f_{n}_{j}.bin'
            if os.path.exists(os.path.join(FAJLY, imya)):
                continue
            zad.append({'task': 'fetch_url', 'ves': ves(imena.get(imya, '')), 'nomer': n,
                        'args': {'url': u.replace('&amp;', '&'), 'insecure': True, 'name': imya}})
    # Сначала полезный разряд, внутри разряда — свежие закупки (номер по убыванию).
    zad.sort(key=lambda z: (z['ves'], -int(z['nomer'])))
    vsego = len(zad)
    zad, otbrosheno = zad[:predel], zad[predel:]
    print(f'вложений к скачиванию: {len(zad)} из {vsego}', file=sys.stderr)
    # Молчаливого обрезания очереди быть не должно: что не поместилось в предел — печатаем.
    if otbrosheno:
        import collections
        c = collections.Counter(z['ves'] for z in otbrosheno)
        print(f'  за пределом осталось {len(otbrosheno)}, по разряду (1=техзадание, '
              f'2=документация, 8=протоколы): {dict(sorted(c.items()))}', file=sys.stderr)
    for z in zad:
        z.pop('ves', None)
        z.pop('nomer', None)
    for i in range(0, len(zad), pachka):
        for r in runner_many(zad[i:i + pachka], pachka):
            d = (r or {}).get('data') or {}
            if d.get('drop_name'):
                skachat(d['drop_name'], FAJLY)
        print(f'  {min(i + pachka, len(zad))}/{len(zad)}', file=sys.stderr, flush=True)



def tekst_iz(p, glubina=0):
    """Текст из файла любого типа. Архив распаковывается и обходится рекурсивно.

    `glubina` — защита от архива в архиве до бесконечности и от zip-бомбы: глубже двух
    уровней не идём.
    """
    # Сигнатуру смотрим по началу файла, а текст берём из ПОЛНЫХ байт. Первая версия читала
    # `read(1_000_000)` и отдавала эти же байты в разбор doc/rtf/html — то есть вносила ровно ту
    # обрезку, которую мы весь день выкорчёвываем.
    with open(p, 'rb') as fh:
        golova = fh.read(4096)
    t = podpis(golova, p)
    if t == 'pdf':
        try:
            import fitz
            with fitz.open(p) as d:
                tekst = '\n'.join(s.get_text() for s in d)
                stranic = len(d)
                # Скан: текстового слоя нет или он с колонтитул. Тот же порог, что у шага
                # `--tekst`. Распознаём страницы, иначе документ уходит в корзину «пусто».
                if len(tekst.strip()) < 200 and stranic and shutil.which('tesseract'):
                    # Страницы распознаются ПАРАЛЛЕЛЬНО: по одной за раз давало 150 файлов за
                    # 15 минут, то есть больше двух часов на корпус при четырёх ядрах.
                    #
                    # Рендер идёт ПОСЛЕДОВАТЕЛЬНО, а параллелится только tesseract, и это не
                    # придирка: PyMuPDF не потокобезопасен на ОДНОМ документе, и первая версия
                    # этой правки дёргала `d[n].get_pixmap()` из потоков — то есть могла отдать
                    # битую картинку или упасть, причём молча и не каждый раз. Узкое место всё
                    # равно tesseract (секунды на страницу), а рендер — миллисекунды.
                    with tempfile.TemporaryDirectory() as tmpd:
                        kartinki = []
                        skolko = min(stranic, PREDEL_STRANIC_OCR, max(0, _bujet['stranic']))
                        if skolko < stranic:
                            _bujet['obrezano'] += stranic - skolko
                            print(f'    OCR обрезан: {os.path.basename(p)} — {stranic} страниц, '
                                  f'распознаю {skolko} (бюджет файла {_bujet["stranic"]})',
                                  file=sys.stderr, flush=True)
                        _bujet['stranic'] -= skolko
                        if skolko == 0:
                            return t, tekst, stranic
                        for i_str in range(skolko):
                            pt = os.path.join(tmpd, f'{i_str:04d}.png')
                            d[i_str].get_pixmap(dpi=200).save(pt)
                            kartinki.append(pt)

                        def _ocr(pt):
                            r = subprocess.run(['tesseract', pt, 'stdout', '-l', 'rus+eng'],
                                               capture_output=True, timeout=300)
                            return r.stdout.decode('utf-8', 'replace')

                        with ThreadPoolExecutor(max_workers=min(6, os.cpu_count() or 2)) as pool:
                            kuski = list(pool.map(_ocr, kartinki))
                    raspoznano = '\n'.join(kuski)
                    if len(raspoznano.strip()) > len(tekst.strip()):
                        return ('pdf-ocr-обрезан' if stranic > PREDEL_STRANIC_OCR else 'pdf-ocr',
                                raspoznano, stranic)
                return t, tekst, stranic
        except Exception as e:  # noqa: BLE001
            return t, f'__ОШИБКА__ {type(e).__name__}: {str(e)[:60]}', 0
    if t in ('docx', 'xlsx', 'ooxml-иное', 'zip/docx/xlsx', 'zip-битый'):
        # Абзацев мало: замер по корпусу — 16 файлов docx отдавали пустой текст, потому что
        # `docx.Document(p).paragraphs` не видит ТАБЛИЦ, а «Главный механик — Иванов И.И. — тел.»
        # в документах закупок сплошь и рядом стоит именно в таблице. Берём абзацы и таблицы, а
        # если вышло пусто — падаем на сырой xml, где есть весь текст, включая надписи.
        try:
            import docx
            d = docx.Document(p)
            chasti = [x.text for x in d.paragraphs]
            for tb in d.tables:
                for row in tb.rows:
                    chasti += [c.text for c in row.cells]
            tekst = '\n'.join(x for x in chasti if x and x.strip())
            if len(tekst.strip()) >= 200:
                return t, tekst, 0
        except Exception:  # noqa: BLE001
            tekst = ''
        try:
            with zipfile.ZipFile(p) as z:
                xml = b' '.join(z.read(n) for n in z.namelist()
                                if n.endswith('.xml'))[:4_000_000]
            syroy = re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', ' ', xml.decode('utf-8', 'replace')))
            return t, syroy if len(syroy.strip()) > len(tekst.strip()) else tekst, 0
        except Exception as e:  # noqa: BLE001
            return t, tekst or f'__ОШИБКА__ {str(e)[:60]}', 0
    if t in ('архив', '7z', 'rar'):
        # Настоящий архив: распаковать во временную папку и обойти вложенные файлы тем же
        # разбором. Прежде этот путь отсутствовал вовсе, и «Документация.zip» на 553 КБ
        # проходила как «пусто» — то есть техзадание, за которым всё и затевалось.
        if glubina >= 2:
            return t, '__ОШИБКА__ архив глубже двух уровней', 0
        kuski, vnutri = [], 0
        with tempfile.TemporaryDirectory() as tmp:
            try:
                if t == 'архив':
                    with zipfile.ZipFile(p) as z:
                        z.extractall(tmp)
                elif t == 'rar':
                    # `7z` из p7zip-full RAR не распаковывает (кодек несвободный) и возвращает
                    # ошибку — по замеру так молча терялись 18 файлов корпуса. `unar` умеет.
                    r = subprocess.run(['unar', '-q', '-f', '-o', tmp, p],
                                       capture_output=True, timeout=600)
                    if r.returncode != 0:
                        return t, f'__ОШИБКА__ unar: {r.stderr[:80].decode("utf-8", "replace")}', 0
                else:
                    r = subprocess.run(['7z', 'x', '-y', f'-o{tmp}', p],
                                       capture_output=True, timeout=300)
                    if r.returncode != 0:
                        return t, f'__ОШИБКА__ 7z: {r.stderr[:80].decode("utf-8", "replace")}', 0
            except Exception as e:  # noqa: BLE001
                return t, f'__ОШИБКА__ распаковка: {type(e).__name__}: {str(e)[:60]}', 0
            for koren, _, fajly in os.walk(tmp):
                for f in sorted(fajly):
                    vnutri += 1
                    vp = os.path.join(koren, f)
                    if os.path.getsize(vp) > 80_000_000:
                        continue
                    _, tk, _ = tekst_iz(vp, glubina + 1)
                    if tk and not tk.startswith('__ОШИБКА__'):
                        kuski.append(f'--- {f}\n{tk}')
        return f'{t}({vnutri})', '\n'.join(kuski), 0
    if t == 'html':
        b = open(p, 'rb').read()
        return t, re.sub(r'<[^>]+>', ' ', b.decode('utf-8', 'replace')), 0
    if t in ('doc/xls', 'rtf'):
        # Без внешних утилит достаём хотя бы кириллицу из потока: этого хватает,
        # чтобы увидеть «Главный механик Иванов И.И.» и телефон.
        s = open(p, 'rb').read().decode('cp1251', 'replace')
        return t, re.sub(r'[^\wа-яёА-ЯЁ.,()@+\-/ \n]', ' ', s), 0
    return t, '', 0


def imena_dokumentov():
    """Имя документа для каждого ожидаемого файла — из тултипа рядом со ссылкой.

    Заодно это и есть карта «файл → закупка», причём детерминированная: номер закупки стоит в
    самом имени файла (`eis_f_<закупка>_<j>.bin`), а `karta.json` от прогона качалки для этого
    не нужна — её отсутствие прежде оставляло колонку `zakupka` пустой, и находки нельзя было
    привязать к закупке.
    """
    tultip = re.compile(
        r'href="(https://zakupki\.gov\.ru/223/filestore/public/1\.0/download/fz223/'
        r'file\.html\?uid=[0-9A-F]+)"[^>]*data-tooltip=\'<span[^>]*>([^<]{1,200})</span>', re.I)
    out = {}
    for p in sorted(glob.glob(os.path.join(DOKI, 'eis_d_*.html'))):
        n = re.search(r'eis_d_(\d+)', p).group(1)
        h = open(p, encoding='utf-8', errors='replace').read()
        po_uid = {}
        for u, im in tultip.findall(h):
            po_uid.setdefault(u.replace('&amp;', '&'), im)
        for j, u in enumerate(dict.fromkeys(SSYLKA_FAJL.findall(h)), 1):
            out[f'eis_f_{n}_{j}.bin'] = po_uid.get(u.replace('&amp;', '&'), '')
    return out


def shag_tekst():
    """Вложения → текст. Пишет тяжёлый дамп с текстом и тонкую сводку без текста.

    Два предела здесь пришлось мерить, а не назначать.

    **Дедуп по содержимому.** ЕИС отдаёт один и тот же файл под разными uid внутри одной
    закупки: из 678 файлов уникальных по содержимому 536, лишних копий 142, а
    «Закупочная_документация.zip» на 12,6 млн знаков лежит в ДЕВЯТИ копиях. Распаковывать и
    разбирать её девять раз — это девять прогонов 7z и девятикратный вес дампа. Текст берётся
    один раз на содержимое, копии ссылаются на первую через `kopiya_ot`.

    **Предел на текст.** Был 2 000 000 знаков, и он резал: 27 файлов упёрлись, настоящих знаков
    у них 257 млн против сохранённых 54 млн. После дедупа предел поднят до 20 млн — этого хватает
    самому длинному файлу корпуса (12,6 млн), а каждый упор считается и печатается.

    **Почему тяжёлый дамп не в git.** С текстом дамп весит 147 МБ и в репозиторий не годится;
    правило проекта — тяжёлое живёт на дропе (`na_drop.sh`). В git идёт `vlozheniya-svodka.csv`:
    те же строки без текста, по ней видно тип, отказ, скан, число знаков и документ.
    """
    os.makedirs(OUT, exist_ok=True)
    imena = imena_dokumentov()
    rows, upor = [], 0
    kesh = {}
    # Строки ПИШУТСЯ СРАЗУ, а не копятся до конца прогона. Прежде весь текст корпуса лежал в
    # памяти (замер на живом прогоне: 672 МБ на середине 1 355 файлов при пределе 20 млн знаков
    # на файл), и смерть на предпоследнем файле стоила бы всего прогона — до конца не
    # записывалось ни строки. Заодно снимается слепота: по размеру jsonl видно, что прогон идёт.
    # Возобновление: уже разобранные файлы пропускаются. Прогон по корпусу идёт часами (OCR
    # сканов и распаковка архивов), и терять его из-за перезапуска нельзя.
    put_js = os.path.join(OUT, 'vlozheniya-tekst.jsonl')
    sdelano = {}
    if os.path.exists(put_js) and '--zanovo' not in sys.argv:
        for l in open(put_js, encoding='utf-8'):
            try:
                x = json.loads(l)
            except Exception:  # noqa: BLE001  оборванная последняя строка после падения
                continue
            sdelano[x['fajl']] = x
            kesh.setdefault(x['sha256'], (x['tip'], '', x['stranic'], x['kopiya_ot'] or x['fajl']))
        rows = [{k: v for k, v in x.items() if k != 'tekst'} for x in sdelano.values()]
        print(f'возобновление: уже разобрано {len(sdelano)} файлов', file=sys.stderr)
    fjs = open(put_js, 'a' if sdelano else 'w', encoding='utf-8')
    vsego = len(glob.glob(os.path.join(FAJLY, 'eis_f_*.bin')))
    for nomer, p in enumerate(sorted(glob.glob(os.path.join(FAJLY, 'eis_f_*.bin'))), 1):
        if os.path.basename(p) in sdelano:
            continue
        imya = os.path.basename(p)
        nachalo = time.time()
        _bujet['stranic'], _bujet['obrezano'] = PREDEL_STRANIC_OCR_FAJL, 0
        if os.path.getsize(p) > 3_000_000:
            print(f'    начат тяжёлый файл {imya} ({os.path.getsize(p) // 1024} КБ)',
                  file=sys.stderr, flush=True)
        h = hashlib.sha256(open(p, 'rb').read()).hexdigest()
        kopiya_ot = ''
        if h in kesh:
            tip, t, stranic, kopiya_ot = (*kesh[h][:3], kesh[h][3])
            if not t and kopiya_ot != imya:
                # текст первоисточника уже на диске; копии он не нужен, она хранит ссылку
                t = ''
        else:
            tip, t, stranic = tekst_iz(p)
            kesh[h] = (tip, t, stranic, imya)
        # Отказ хранилища — НЕ пустота и НЕ скан. Пока они в одной корзине, не видно ни
        # сколько документов надо распознавать, ни сколько перезапрашивать.
        otkaz = tip == 'отказ'
        # Порог по длине, а не «если пусто»: скан с колонтитулом даёт короткий текстовый слой.
        skan = not otkaz and len(t.strip()) < 200
        m = re.search(r'eis_f_(\d+)_', imya)
        if len(t) > PREDEL_TEKSTA:
            upor += 1
        rows.append({'fajl': imya, 'zakupka': m.group(1) if m else '',
                     'dokument': imena.get(imya, ''), 'tip': tip,
                     'bajt': os.path.getsize(p), 'stranic': stranic,
                     'znakov': len(t.strip()), 'otkaz': '1' if otkaz else '',
                     'skan_ili_pusto': '1' if skan else '',
                     'sha256': h[:16], 'kopiya_ot': kopiya_ot if kopiya_ot != imya else '',
                     'ocr_stranic_obrezano': _bujet['obrezano'] or '',
                     # У копии текст не повторяется: с повтором дамп раздувался с 147 до 390 МБ
                     # (142 копии из 678, одна документация в девяти экземплярах). Текст берётся
                     # по `sha256` из строки-первоисточника — так делает shag_lica.
                     'tekst': '' if kopiya_ot and kopiya_ot != imya else t[:PREDEL_TEKSTA]})
        fjs.write(json.dumps(rows[-1], ensure_ascii=False) + '\n')
        fjs.flush()
        # Тяжёлый текст в памяти не держим: он уже на диске, дальше нужны только счётчики.
        rows[-1] = {k: v for k, v in rows[-1].items() if k != 'tekst'}
        # Долгий файл называется вслух. Без этого прогон стоит вслепую: один архив на 36 МБ с
        # десятками сканов внутри держит очередь двадцать минут, и снаружи это неотличимо от
        # зависания. Замер живого прогона: `213066_ЗД_ЭФ.rar`, 22 страницы OCR за 13 минут.
        if time.time() - nachalo > 60:
            print(f'  долгий файл {imya} ({os.path.getsize(p)//1024} КБ, {tip}): '
                  f'{time.time() - nachalo:.0f} с', file=sys.stderr, flush=True)
        if nomer % 50 == 0:
            print(f'  {nomer}/{vsego}: с текстом '
                  f'{sum(1 for x in rows if not (x["skan_ili_pusto"] or x["otkaz"]))}, '
                  f'отказов {sum(1 for x in rows if x["otkaz"])}, '
                  f'уникальных содержимым {len(kesh)}', file=sys.stderr, flush=True)
    fjs.close()
    # Тонкая сводка без текста — она и уходит в git, по ней виден весь корпус одним взглядом.
    svodka = os.path.join(OUT, 'vlozheniya-svodka.csv')
    kol = ['fajl', 'zakupka', 'dokument', 'tip', 'bajt', 'stranic', 'znakov', 'otkaz',
           'skan_ili_pusto', 'sha256', 'kopiya_ot', 'ocr_stranic_obrezano']
    with open(svodka, 'w', encoding='utf-8-sig', newline='') as f:
        w = csv.DictWriter(f, fieldnames=kol, delimiter=';', extrasaction='ignore')
        w.writeheader()
        for r in rows:
            w.writerow(r)
    import collections
    print(f'файлов: {len(rows)}')
    print('  по типу:', dict(collections.Counter(r['tip'] for r in rows)))
    print(f'  отказ хранилища («не опубликован»): {sum(1 for r in rows if r["otkaz"])}')
    print(f'  скан или пусто (<200 знаков), отказы не считая: '
          f'{sum(1 for r in rows if r["skan_ili_pusto"])}')
    print(f'  с текстом: {sum(1 for r in rows if not (r["skan_ili_pusto"] or r["otkaz"]))}')
    print(f'  уникальных по содержимому: {len(kesh)}, копий использовано повторно: '
          f'{len(rows) - len(kesh)}')
    if upor:
        print(f'  ВНИМАНИЕ: текст упёрся в предел {PREDEL_TEKSTA} знаков у {upor} файлов')
    # Быстрый механический замер: есть ли вообще пары «должность + ФИО»
    par, tel = 0, 0
    for r in (json.loads(l) for l in open(os.path.join(OUT, 'vlozheniya-tekst.jsonl'),
                                          encoding='utf-8')):
        if r['skan_ili_pusto'] or r['otkaz'] or not r.get('tekst'):
            continue
        t = r['tekst']
        inn = {cifry(x) for x in INN_OGRN.findall(t)}
        for m in DOLZH.finditer(t):
            # Окно СИММЕТРИЧНОЕ, и это исправление по контролю на эталонах. Прежде окно шло
            # только вперёд от должности, и запись «Александр Владимирович Лосев, Заместитель
            # главного механика» не находилась вовсе: имя стоит ПЕРЕД должностью. В документах
            # встречаются оба порядка — «УТВЕРЖДАЮ Главный инженер Иванов И.И.» и «Иванов И.И.,
            # главный инженер», — поэтому смотреть надо в обе стороны.
            okno = t[max(0, m.start() - 250):m.start() + 250]
            if FIO.search(okno):
                par += 1
                for c in SVYAZ.findall(okno):
                    if cifry(c) not in inn:
                        tel += 1
    print(f'  пар «должность рядом с ФИО»: {par}, из них с телефоном рядом: {tel}')
    print(f'→ {os.path.join(OUT, "vlozheniya-tekst.jsonl")}')


def shag_lica(threads=8, pachka=3):
    import gen_provider as G
    G.env = lambda: {'PROVIDER_API_KEY': os.environ['PROVIDER_API_KEY'],
                     'PROVIDER_BASE_URL': os.environ.get('PROVIDER_BASE_URL',
                                                         'https://router.cheap')}
    src = os.path.join(OUT, 'vlozheniya-tekst.jsonl')
    rows = [json.loads(l) for l in open(src, encoding='utf-8')]
    # Текст копии лежит в строке-первоисточнике, находим его по sha256.
    po_sha = {r['sha256']: r['tekst'] for r in rows if r.get('tekst')}
    for r in rows:
        if not r.get('tekst') and r.get('kopiya_ot'):
            r['tekst'] = po_sha.get(r['sha256'], '')
    # Один и тот же документ в одной и той же закупке спрашивать дважды незачем; в РАЗНОЙ
    # закупке — нужно, там другой заказчик и другие люди рядом.
    vidno = set()
    # Провайдеру отдаём только куски вокруг должностей: целиком техзадание это сотни страниц
    # труб и болтов, а нам нужны подписи и контакты. Это же бережёт баланс.
    # Пределы здесь были 8 кусков и 9 000 знаков на файл, и оба резали по живому: замер по
    # `vlozheniya-tekst.jsonl` показал 14 файлов из 32 с БОЛЬШЕ чем 8 упоминаниями должности
    # (в одном 71) и 4 файла со склейкой длиннее 9 000. То есть в самых густых документах —
    # там, где люди и есть, — до провайдера доходила пятая часть упоминаний.
    #
    # Кроме подъёма предела окна СЛИВАЮТСЯ: при 71 упоминании окна ±300/+400 знаков сплошь
    # перекрываются, и без слияния один и тот же текст уезжает в запрос по многу раз. Слияние
    # и полнее, и дешевле по балансу.
    PREDEL_OKON, PREDEL_ZNAKOV = 40, 40000
    zadaniya, upor_okna, upor_znaki = [], 0, 0
    for r in rows:
        if r['skan_ili_pusto'] or r['otkaz']:
            continue
        t = r['tekst']
        klyuch = (r['zakupka'], r['sha256'])
        if klyuch in vidno:
            continue
        vidno.add(klyuch)
        rangi = [(max(0, m.start() - 300), m.start() + 400) for m in DOLZH.finditer(t)]
        if not rangi:
            continue
        slito = []
        for a, b in rangi:
            if slito and a <= slito[-1][1]:
                slito[-1] = (slito[-1][0], max(slito[-1][1], b))
            else:
                slito.append((a, b))
        if len(slito) > PREDEL_OKON:
            upor_okna += 1
            slito = slito[:PREDEL_OKON]
        telo = '\n---\n'.join(t[a:b] for a, b in slito)
        if len(telo) > PREDEL_ZNAKOV:
            upor_znaki += 1
            telo = telo[:PREDEL_ZNAKOV]
        zadaniya.append((r['fajl'], r['zakupka'], telo))
    print(f'файлов с должностями: {len(zadaniya)}', file=sys.stderr)
    if upor_okna or upor_znaki:
        print(f'  упёрлись в предел: окон {upor_okna}, знаков {upor_znaki} '
              f'(пределы {PREDEL_OKON} и {PREDEL_ZNAKOV})', file=sys.stderr)

    PROMPT = """Ты читаешь куски документов закупки промышленного предприятия (техзадание,
извещение, протокол, лист согласования). Нужны люди в ТЕХНИЧЕСКИХ ролях: главный инженер,
главный механик, главный энергетик, начальник цеха или компрессорной, их заместители, инженер
ОГМ/ОГЭ/КИПиА, технический директор. Снабженец помечается отдельно и стоит ниже.

ПРАВИЛА, нарушать нельзя:
1. **Ни одного телефона и ни одной фамилии, которых нет во входном тексте.** Выдуманный
   контакт хуже пропуска: по нему позвонят чужому человеку.
2. Число из десяти цифр — не обязательно телефон. Если рядом стоит «ИНН», «ОГРН», «р/с»,
   «КПП» или это номер документа — это не телефон, не пиши его.
3. Не считай технической ролью «главный инженер проекта» и «ГИП»: это проектировщик, машину
   на площадке он не выбирает. Также не считай «механик гаража», «начальник транспортного
   цеха», «инженер по закупкам», «инженер-сметчик».
4. Если людей в тексте нет — верни пустой список. Не угадывай.

ОТВЕТ — строго JSON, без пояснений:
{"lyudi":[{"imya":"","dolzhnost":"","rol":"техническая|снабжение|неясно","telefon":"",
"pochta":"","citata":"кусок текста, где это написано, до 150 знаков"}]}

ТЕКСТ:
"""
    lock = threading.Lock()
    client = G.make_client()
    itog, sch = [], {'lyudey': 0, 'teh': 0, 'vydumka': 0, 'err': 0}

    def odin(z):
        fajl, zak, telo = z
        try:
            o = G.call(client, [{'role': 'user', 'content': PROMPT + telo}],
                       model='claude-fable-5', attempts=4)
            txt = ''.join(b.text for b in o.content if b.type == 'text')
            m = re.search(r'\{.*\}', txt, re.S)
            return fajl, zak, telo, (json.loads(m.group(0)) if m else None)
        except Exception as e:  # noqa: BLE001
            return fajl, zak, telo, {'__err': f'{type(e).__name__}: {str(e)[:60]}'}

    with ThreadPoolExecutor(max_workers=threads) as pool:
        for fajl, zak, telo, res in pool.map(odin, zadaniya):
            with lock:
                if not res or '__err' in (res or {}):
                    sch['err'] += 1
                    continue
                cif_vhod = cifry(telo)
                for ch in res.get('lyudi') or []:
                    t = ch.get('telefon') or ''
                    est = (not t) or (cifry(t)[-10:] and cifry(t)[-10:] in cif_vhod)
                    if t and not est:
                        sch['vydumka'] += 1
                    itog.append({**ch, 'fajl': fajl, 'zakupka': zak,
                                 'telefon_est_v_tekste': '1' if est else ''})
                    sch['lyudey'] += 1
                    if ch.get('rol') == 'техническая':
                        sch['teh'] += 1

    out = os.path.join(OUT, 'vlozheniya-lica.csv')
    cols = ['zakupka', 'fajl', 'imya', 'dolzhnost', 'rol', 'telefon', 'pochta',
            'telefon_est_v_tekste', 'citata']
    with open(out, 'w', encoding='utf-8-sig', newline='') as f:
        w = csv.DictWriter(f, fieldnames=cols, delimiter=';', extrasaction='ignore')
        w.writeheader()
        for r in itog:
            w.writerow(r)
    print(f'людей: {sch["lyudey"]}, технических: {sch["teh"]}, '
          f'номеров не из текста: {sch["vydumka"]}, сбоев: {sch["err"]} → {out}', file=sys.stderr)


if __name__ == '__main__':
    if '--doki' in sys.argv:
        shag_doki()
    elif '--fajly' in sys.argv:
        shag_fajly(predel=int(sys.argv[sys.argv.index('--predel') + 1])
                   if '--predel' in sys.argv else 200)
    elif '--tekst' in sys.argv:
        shag_tekst()
    elif '--lica' in sys.argv:
        shag_lica()
    else:
        sys.exit(__doc__)
